# Pull the data from the Google Drive and Move it to the folder 'New Data'


import os
import shutil
import sqlite3
import xml.etree.ElementTree as ET
import re
import scipy.stats as stats
import matplotlib.pyplot as plt
from docx import Document
from matplotlib.backends.backend_pdf import PdfPages
import datetime
import shutil
import tkinter as tk
from tkinter import filedialog


def select_folder():
    root = tk.Tk()
    root.withdraw()  # Hide the main window

    folder_path = filedialog.askdirectory(initialdir=r"H:\Pitching\Data", title="Select Folder")
    if folder_path:
        copy_and_move_files(folder_path)


def extract_initials(folder_name):
    initials = ''.join(filter(str.isupper, folder_name))
    return initials[::-1]  # Reverse the initials


def copy_and_move_files(base_path):
    # Get the base folder name
    base_folder = os.path.basename(base_path)

    # Get list of folders in the base path
    folders = [f for f in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, f))]
    # Sort folders by creation time
    folders.sort(key=lambda x: os.path.getctime(os.path.join(base_path, x)), reverse=True)

    if not folders:
        print("No folders found in the specified directory.")
        return

    latest_folder = folders[0]
    folder_path = os.path.join(base_path, latest_folder)

    # Get creation time of the latest folder
    creation_time = os.path.getctime(folder_path)
    creation_date = datetime.datetime.fromtimestamp(creation_time).strftime('%m-%d-%y')

    # Find the files to copy and rename
    files_to_copy_rename = []
    for file in os.listdir(folder_path):
        if file.lower() == 'session.xml' or file.lower() == 'session_data.xml':
            files_to_copy_rename.append(file)

    if not files_to_copy_rename:
        print("No XML files named 'session.xml' or 'session_data.xml' found in the latest folder.")
        return

    # Construct new file names
    new_file_names = []
    initials = extract_initials(base_folder)
    for file in files_to_copy_rename:
        if file.lower() == 'session.xml':
            new_file_name = f"session_{initials}_{creation_date}.xml"
        elif file.lower() == 'session_data.xml':
            new_file_name = f"session_data_{initials}_{creation_date}.xml"
        else:
            continue
        new_file_names.append(new_file_name)

    # Create temporary directory to store renamed copies
    temp_dir = os.path.join(base_path, "temp_rename")
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)

    # Copy and rename the files to the temporary directory
    for original_file, new_file_name in zip(files_to_copy_rename, new_file_names):
        original_file_path = os.path.join(folder_path, original_file)
        new_file_path = os.path.join(temp_dir, new_file_name)
        shutil.copyfile(original_file_path, new_file_path)
        print(f"File '{original_file}' copied and renamed to: {new_file_path}")

    # Copy the renamed files to the destination folder
    new_folder_path = r"F:\My Drive\Pitching Data\New Data"
    if not os.path.exists(new_folder_path):
        os.makedirs(new_folder_path)
    for new_file_name in new_file_names:
        temp_file_path = os.path.join(temp_dir, new_file_name)
        destination_file_path = os.path.join(new_folder_path, new_file_name)
        shutil.copyfile(temp_file_path, destination_file_path)
        print(f"File '{new_file_name}' copied to: {destination_file_path}")

    # Clean up temporary directory
    shutil.rmtree(temp_dir)


# Prompt user to select the folder
select_folder()


# Source and destination directory paths
source_dir = r"F:\My Drive\Pitching Data\New Data"
destination_dir = r"C:\Users\q\PycharmProjects\Rename_Session.xml\Pitching Data\New Data"

# Iterate over files in the source directory
for file_name in os.listdir(source_dir):
    if file_name.endswith('.xml'):
        source_file_path = os.path.join(source_dir, file_name)
        destination_file_path = os.path.join(destination_dir, file_name)
        shutil.move(source_file_path, destination_file_path)
        print(f"File '{file_name}' moved to: {destination_file_path}")

print("XML files copied successfully!")

# This code is *chef's kiss* perfect. Does everything for parsing for variables, iterates over it
# multiple times, and includes the specific filename for each trial. I love it.


# Connect to the SQLite database
conn = sqlite3.connect("grading_equation_new.db")
cursor = conn.cursor()

# Drop the table if it exists
cursor.execute('DROP TABLE IF EXISTS variables')

# Create the variables table
cursor.execute('''
   CREATE TABLE variables (
       id INTEGER PRIMARY KEY AUTOINCREMENT,
       session_data_id TEXT NOT NULL,
       file_name TEXT NOT NULL,
       Pitch TEXT,
       Score REAL,
       Linear_Pelvis_Speed REAL,
       HSS_Footplant REAL,
       Pelvis_Ang_Footplant REAL,
       Trunk_Ang_Footplant REAL,
       Pelvic_Obl REAL,
       Front_Leg_Brace REAL,
       Front_Leg_Var_Val REAL,
       Lead_Leg_Midpoint REAL,
       Lead_Leg_GRF_y REAL,
       Lead_Leg_GRF_z REAL,
       Lead_Leg_GRF_x REAL,
       Horizontal_Abduction REAL,
       Shld_ER_Footplant REAL,
       Shld_ER_Max REAL,
       Lateral_Trunk_Tilt REAL,
       Pelvis_Ang_Velo REAL,
       Torso_Ang_Velo REAL,
       Arm_Ang_Velo REAL,
       MPH REAL
   )
''')

# Commit the changes before proceeding to parsing XML files
conn.commit()

# List of XML files to parse (replace this with your actual directory path)
directory_path = r"C:\Users\q\PycharmProjects\Rename_Session.xml\Pitching Data\New Data"

while True:
    # Get a list of XML files in the specified directory
    xml_files = [os.path.join(directory_path, file) for file in os.listdir(directory_path) if file.endswith(".xml")]

    # Track the number of new data entries
    new_data_entries = 0

    for file_name in xml_files:
        try:
            # Check if the file has already been processed
            cursor.execute('SELECT COUNT(*) FROM variables WHERE file_name = ?', (file_name,))
            count = cursor.fetchone()[0]

            if count == 0:
                # Parse the XML file
                tree = ET.parse(file_name)
                root = tree.getroot()

                # Extract session_data_id from the file name
                session_data_id_match = re.search(r"session_data_(.+)\.xml", file_name)
                if session_data_id_match:
                    session_data_id = session_data_id_match.group(1)
                else:
                    session_data_id = None

                # Check if session_data_id is None, and skip the file if it is
                if session_data_id is not None:
                    # Iterate over all occurrences of "owner" element with a value containing "Fastball"
                    for owner_element in root.findall('.//owner'):
                        if 'Fastball' in owner_element.get('value', ''):
                            fastball_owner = owner_element
                            # Extract and insert data into the database
                            linear_pelvis_speed = None
                            hss_footplant = None
                            pelvis_ang_fp = None
                            trunk_ang_fp = None
                            pelvis_obl = None
                            front_leg_brace = None
                            front_leg_var_val = None
                            lead_leg_midpoint = None
                            lead_leg_grf_y = None
                            lead_leg_grf_z = None
                            lead_leg_grf_x = None
                            horizontal_abduction = None
                            shld_er_fp = None
                            shld_er_max = None
                            lateral_trunk_tilt = None
                            pelvis_ang_velo = None
                            torso_ang_velo = None
                            arm_ang_velo = None

                            # Extract the value of "owner"
                            pitch_value = fastball_owner.get('value', '')

                            for variable_element in fastball_owner.findall('.//name[@value]'):
                                variable_name = variable_element.attrib['value']
                                component_x_element = variable_element.find('./component[@value="X"]')
                                component_y_element = variable_element.find('./component[@value="Y"]')
                                component_z_element = variable_element.find('./component[@value="Z"]')

                                if variable_name == "MaxPelvisLinearVel_MPH" and component_y_element is not None:
                                    linear_pelvis_speed = float(component_y_element.attrib['data'].replace(',', '.'))
                                elif variable_name == "Hip Shoulders Sep@Footstrike" and component_z_element is not None:
                                    hss_footplant = float(component_z_element.attrib['data'].replace(',', '.'))
                                elif variable_name == "Pelvis_Angle@Footstrike" and component_z_element is not None:
                                    pelvis_ang_fp = float(component_z_element.attrib['data'].replace(',', '.'))
                                elif variable_name == "Trunk_Angle@Footstrike" and component_z_element is not None:
                                    trunk_ang_fp = float(component_z_element.attrib['data'].replace(',', '.'))
                                elif variable_name == "Lead_Knee_Angle@Footstrike" and component_x_element is not None:
                                    lead_knee_ang_at_footstrike = float(
                                        component_x_element.attrib['data'].replace(',', '.'))
                                    lead_knee_ang_at_release = float(root.find(
                                        './/name[@value="Lead_Knee_Angle@Release"]/component[@value="X"]').attrib[
                                                                         'data'].replace(',', '.'))
                                    front_leg_brace = lead_knee_ang_at_footstrike - lead_knee_ang_at_release
                                if variable_name == "Lead_Knee_Angle@Footstrike" and component_y_element is not None:
                                    lead_knee_var_val_fp = float(component_y_element.attrib['data'].replace(',', '.'))
                                elif variable_name == "Lead_Knee_Angle@Release" and component_y_element is not None:
                                    lead_knee_var_val_rel = float(component_y_element.attrib['data'].replace(',', '.'))
                                    # Calculate the difference and store in front_leg_var_val
                                    front_leg_var_val = lead_knee_var_val_fp - lead_knee_var_val_rel
                                if variable_name == "Pelvis_Angle@Footstrike" and component_z_element is not None:
                                    pelvis_obl_fp = float(component_y_element.attrib['data'].replace(',', '.'))
                                    pelvis_angle_release_element = root.find(
                                        './/name[@value="Pelvis_Angle@Release"]/component[@value="Y"]')
                                    if pelvis_angle_release_element is not None:
                                        pelvis_obl_release = float(
                                            pelvis_angle_release_element.attrib['data'].replace(',', '.'))
                                        pelvis_obl = pelvis_obl_release - pelvis_obl_fp
                                elif variable_name == "Lead_Leg_GRF_mag_Midpoint_FS_Release" and component_x_element is not None:
                                    lead_leg_midpoint = abs(float(component_x_element.attrib['data'].replace(',', '.')))
                                elif variable_name == "Lead_Leg_GRF_min" and component_y_element is not None:
                                    lead_leg_grf_y = abs(float(component_y_element.attrib['data'].replace(',', '.')))
                                elif variable_name == "Lead_Leg_GRF_max" and component_z_element is not None:
                                    lead_leg_grf_z = float(component_z_element.attrib['data'].replace(',', '.'))
                                if variable_name == "Lead_Leg_GRF_max" and component_x_element is not None:
                                    lat_force_max = float(component_x_element.attrib['data'].replace(',', '.'))
                                elif variable_name == "Lead_Leg_GRF_min" and component_x_element is not None:
                                    lat_force_min = float(component_x_element.attrib['data'].replace(',', '.'))
                                    # Calculate the sum of absolute values and store in lead_leg_grf_x
                                    lead_leg_grf_x = abs(lat_force_max) + abs(lat_force_min)
                                elif variable_name == "Pitching_Shoulder_Angle@Footstrike" and component_x_element is not None:
                                    horizontal_abduction = abs(
                                        float(component_x_element.attrib['data'].replace(',', '.')))
                                if variable_name == "Pitching_Shoulder_Angle@Footstrike":
                                    if component_z_element is not None:
                                        shld_er_fp = float(component_z_element.attrib['data'].replace(',', '.'))
                                elif variable_name == "Pitching_Shoulder_Angle_Max" and component_z_element is not None:
                                    shld_er_max = abs(float(component_z_element.attrib['data'].replace(',', '.')))
                                elif variable_name == "Trunk_wrt_Pelvis_FE@Release" and component_y_element is not None:
                                    lateral_trunk_tilt = float(component_y_element.attrib['data'].replace(',', '.'))
                                elif variable_name == "Pelvis_Ang_Vel_max" and component_x_element is not None:
                                    pelvis_ang_velo = float(component_x_element.attrib['data'].replace(',', '.'))
                                elif variable_name == "Thorax_Ang_Vel_max" and component_x_element is not None:
                                    torso_ang_velo = float(component_x_element.attrib['data'].replace(',', '.'))
                                elif variable_name == "Pitching_Humerus_Ang_Vel_max" and component_x_element is not None:
                                    arm_ang_velo = float(component_x_element.attrib['data'].replace(',', '.'))

                                linear_pelvis_speed = linear_pelvis_speed if linear_pelvis_speed is not None else 4
                                hss_footplant = hss_footplant if hss_footplant is not None else 30
                                pelvis_ang_fp = pelvis_ang_fp if pelvis_ang_fp is not None else 60
                                trunk_ang_fp = trunk_ang_fp if trunk_ang_fp is not None else 90
                                pelvis_obl = pelvis_obl if pelvis_obl is not None else 2
                                front_leg_brace = front_leg_brace if front_leg_brace is not None else 8
                                front_leg_var_val = front_leg_var_val if front_leg_var_val is not None else 0
                                lead_leg_midpoint = lead_leg_midpoint if lead_leg_midpoint is not None else 1.5
                                lead_leg_grf_y = lead_leg_grf_y if lead_leg_grf_y is not None else .7
                                lead_leg_grf_z = lead_leg_grf_z if lead_leg_grf_z is not None else 1.5
                                lead_leg_grf_x = lead_leg_grf_x if lead_leg_grf_x is not None else .25
                                horizontal_abduction = horizontal_abduction if horizontal_abduction is not None else 20
                                shld_er_fp = shld_er_fp if shld_er_fp is not None else 45
                                lateral_trunk_tilt = lateral_trunk_tilt if lateral_trunk_tilt is not None else 35
                                shld_er_max = shld_er_max if shld_er_max is not None else 160
                                pelvis_ang_velo = pelvis_ang_velo if pelvis_ang_velo is not None else 600
                                torso_ang_velo = torso_ang_velo if torso_ang_velo is not None else 1000
                                arm_ang_velo = arm_ang_velo if arm_ang_velo is not None else 5000

                                # Insert the data into the database
                            cursor.execute('''
                               INSERT INTO variables (
                                   session_data_id,
                                   file_name,
                                   Pitch,
                                   Linear_Pelvis_Speed,
                                   HSS_Footplant,
                                   Pelvis_Ang_Footplant,
                                   Trunk_Ang_Footplant,
                                   Pelvic_Obl,
                                   Front_Leg_Brace,
                                   Front_Leg_Var_Val,
                                   Lead_Leg_Midpoint,
                                   Lead_Leg_GRF_y,
                                   Lead_Leg_GRF_z,
                                   Lead_Leg_GRF_x,
                                   Horizontal_Abduction,
                                   Shld_ER_Footplant,
                                   Shld_ER_Max,
                                   Lateral_Trunk_Tilt,
                                   Pelvis_Ang_Velo,
                                   Torso_Ang_Velo,
                                   Arm_Ang_Velo,
                                   MPH
                               ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                           ''', (
                            session_data_id, file_name, pitch_value, linear_pelvis_speed, hss_footplant, pelvis_ang_fp,
                            trunk_ang_fp, pelvis_obl, front_leg_brace, front_leg_var_val, lead_leg_midpoint, lead_leg_grf_y,
                            lead_leg_grf_z, lead_leg_grf_x, horizontal_abduction, shld_er_fp, shld_er_max,
                            lateral_trunk_tilt, pelvis_ang_velo, torso_ang_velo, arm_ang_velo, None))

                            # Increment the count of new data entries
                            new_data_entries += 1
                else:
                    print(f"Skipping file {file_name} due to missing session_data_id.")

        except Exception as e:
            # Print the error and the file name where it occurred
            print(f"Error processing file {file_name}: {e}")

    # Commit the changes for each file
    conn.commit()

    # If no new data entries, break out of the loop
    if new_data_entries == 0:
        break

# Close the connection
conn.close()

print("Data inserted into the database.")

# Successfully adds MPH to the appropriate spot in the database


# Connect to the database
conn = sqlite3.connect("grading_equation_new.db")
cursor = conn.cursor()

directory_path = r"C:\Users\q\PycharmProjects\Rename_Session.xml\Pitching Data\New Data"

# Get a list of XML files in the specified directory
xml_files = [os.path.join(directory_path, file) for file in os.listdir(directory_path) if file.endswith(".xml")]

# Track the number of new data entries
new_data_entries = 0

for file_name in xml_files:
    try:
        # Check if the file has already been processed
        cursor.execute('SELECT COUNT(*) FROM variables WHERE file_name = ?', (file_name,))
        count = cursor.fetchone()[0]

        if count == 0:
            # Parse the XML file
            tree = ET.parse(file_name)
            root = tree.getroot()

            # Extract session_data_id from the file name
            session_data_id_match = re.search(r"session_(.+)\.xml", file_name)
            if session_data_id_match:
                session_data_id = session_data_id_match.group(1)
            else:
                session_data_id = None

            # Check if session_data_id is None, and skip the file if it is
            if session_data_id is not None:
                # Iterate over all occurrences of "Measurement" element
                for measurement_element in root.findall('.//Measurement'):
                    pitch_value = measurement_element.get('Filename')
                    used_element = measurement_element.find('./Fields/Used')

                    if pitch_value is not None and used_element is not None:
                        # Replace "qtm" with "c3d" in pitch_value
                        pitch_value = pitch_value.replace("qtm", "c3d")

                        # Print statement for debugging
                        print(f"Filename: {pitch_value}, Used: {used_element.text}")

                        if 'Fastball' in pitch_value and used_element.text == 'True':
                            # Extract the values of "Comments" and "Ball_speed"
                            comments_value_element = measurement_element.find('./Fields/Comments')

                            if comments_value_element is not None:
                                print(f"MPH: {comments_value_element.text}")
                                print(f"Session ID: {session_data_id}")

                                mph_value = float(comments_value_element.text)

                                # Update the MPH value in the "variables" table
                                print(f"Updating MPH for session_data_id: {session_data_id}, Pitch: {pitch_value}")
                                cursor.execute('''
                                   UPDATE variables
                                   SET MPH = ?
                                   WHERE session_data_id = ? AND Pitch = ?
                                ''', (mph_value, session_data_id, pitch_value))

                                # Increment the new data entries counter
                                new_data_entries += 1
                    else:
                        print("Filename or Used element not found.")

    except Exception as e:
        print(f"Error processing file {file_name}: {e}")

# Commit the changes
conn.commit()

# Close the connection
conn.close()

print(f"Number of new data entries: {new_data_entries}")
print("Data inserted into the database.")

# Generates Score


# Connect to the SQLite database
conn = sqlite3.connect("grading_equation_new.db")
cursor = conn.cursor()

# Select all rows from the 'variables' table
cursor.execute('SELECT * FROM variables')
rows = cursor.fetchall()

# Iterate through each row and calculate the score
for row in rows:
    linear_pelvis_speed = row[5]
    hss_footplant = row[6]
    pelvis_ang_fp = row[7]
    trunk_ang_fp = row[8]
    pelvis_obl = row[9]
    front_leg_brace = row[10]
    front_leg_var_val = row[11]
    lead_leg_midpoint = row[12]
    lead_leg_grf_y = row[13]
    lead_leg_grf_z = row[14]
    lead_leg_grf_x = row[15]
    horizontal_abduction = row[16]
    shld_er_fp = row[17]
    shld_er_max = row[18]
    lateral_trunk_tilt = row[19]
    pelvis_ang_velo = row[20]
    torso_ang_velo = row[21]
    arm_ang_velo = row[22]
    MPH = row[23]

    # Check for None values and replace with 0
    linear_pelvis_speed = linear_pelvis_speed if linear_pelvis_speed is not None else 4
    hss_footplant = hss_footplant if hss_footplant is not None else 30
    pelvis_ang_fp = pelvis_ang_fp if pelvis_ang_fp is not None else 60
    trunk_ang_fp = trunk_ang_fp if trunk_ang_fp is not None else 90
    pelvis_obl = pelvis_obl if pelvis_obl is not None else 2
    front_leg_brace = front_leg_brace if front_leg_brace is not None else 8
    front_leg_var_val = front_leg_var_val if front_leg_var_val is not None else 0
    lead_leg_midpoint = lead_leg_midpoint if lead_leg_midpoint is not None else 1.5
    lead_leg_grf_y = lead_leg_grf_y if lead_leg_grf_y is not None else .7
    lead_leg_grf_z = lead_leg_grf_z if lead_leg_grf_z is not None else 1.5
    lead_leg_grf_x = lead_leg_grf_x if lead_leg_grf_x is not None else .25
    horizontal_abduction = horizontal_abduction if horizontal_abduction is not None else 20
    shld_er_fp = shld_er_fp if shld_er_fp is not None else 45
    lateral_trunk_tilt = lateral_trunk_tilt if lateral_trunk_tilt is not None else 35
    shld_er_max = shld_er_max if shld_er_max is not None else 160
    pelvis_ang_velo = pelvis_ang_velo if pelvis_ang_velo is not None else 600
    torso_ang_velo = torso_ang_velo if torso_ang_velo is not None else 1000
    arm_ang_velo = arm_ang_velo if arm_ang_velo is not None else 5000
    MPH = MPH if MPH is not None else 75

    # Calculate the score using the specified equation
    score = (0.05 * linear_pelvis_speed) + (0.15 * front_leg_brace) + \
            (5 * lead_leg_midpoint) + \
            (0.55 * horizontal_abduction) + (0.012 * torso_ang_velo) + (.05 * pelvis_obl) + (
                        .6 * trunk_ang_fp) - (.45 * pelvis_ang_fp) + (.1 * shld_er_max) + (.1 * front_leg_var_val) + (
                        .05 * pelvis_ang_velo) + MPH

    # Update the 'Score' column in the database
    cursor.execute('UPDATE variables SET Score = ? WHERE id = ?', (score, row[0]))

# Commit the changes and close the connection
conn.commit()
conn.close()

print("Scores updated in the database.")

# Generates percentile graphs and exports both and PDF and Word Doc for the person
from datetime import date
import os
import sqlite3
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.dml import MSO_THEME_COLOR
from matplotlib import pyplot as plt
from scipy import stats
import docx2txt
from docx import Document

# Connect to the SQLite databases
conn_new = sqlite3.connect("grading_equation_new.db")
conn_reference = sqlite3.connect("grading_equation_reference_data_Pro.db")
cursor_new = conn_new.cursor()
cursor_reference = conn_reference.cursor()

# Define the directory to save exported graphs
export_dir = r"F:\My Drive\Percentile Reports\Pro"

try:
    # Get the session_data_id for naming the Word document
    cursor_new.execute('SELECT session_data_id FROM variables LIMIT 1')
    session_data_id = cursor_new.fetchone()[0]

    # Define the filename for the Word document
    docx_filename = os.path.join(export_dir, f"8ctane Percentiles {session_data_id}.docx")

    # Create a Word document to save the plots
    doc = Document()

    # Save the document as DOCX
    doc.save("output.docx")

    # Convert DOCX to HTML
    html_content = docx2txt.process("output.docx")
    html_filename = "output.html"
    with open(html_filename, "w") as html_file:
        html_file.write(html_content)

    # Define CSS to set the background color
    css = "body { background-color: #393939; color: #FFFFFF; }"

    # # Convert HTML to PDF with background color using WeasyPrint
    # pdf_filename = "output.pdf"
    # HTML(html_filename).write_pdf(pdf_filename, stylesheets=[css])

    print("PDF with background color generated successfully.")

    # Set background color
    background_color = RGBColor(57, 57, 57)  # Dark gray background
    for section in doc.sections:
        section.page_background_color = background_color

    # Set font color
    font_color = RGBColor(255, 255, 255)  # White font color
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = font_color

    # Add image "8ctane - Faded 8 to Blue.png" at the top center of the page
    image_filename = "8ctane - Faded 8 to Blue.png"
    image_path = os.path.join(export_dir, image_filename)
    doc.add_picture(image_path, width=Inches(4.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add player's name and date
    doc.add_paragraph("\nPlayers Name: ")
    doc.add_paragraph(f"Date: {date.today().strftime('%B %d, %Y')}")

    # Get the list of columns in the new database
    cursor_new.execute("PRAGMA table_info(variables)")
    column_names = [row[1] for row in cursor_new.fetchall() if
                    row[1] not in ('id', 'session_data_id', 'file_name', 'Pitch')]

    # Retrieve the row with the highest score from the new database
    cursor_new.execute('SELECT * FROM variables ORDER BY Score DESC LIMIT 1')
    highest_score_row = cursor_new.fetchone()

    # Get the values from the highest scoring row
    highest_score_values = highest_score_row[4:]  # Exclude 'id', 'session_data_id', 'file_name', and 'Pitch'

    # Determine the layout for each histogram
    histogram_layout = {
        "Score": {"width": Inches(6.0), "alignment": WD_ALIGN_PARAGRAPH.CENTER},
        "Linear_Pelvis_Speed": {"width": Inches(3.0), "alignment": WD_ALIGN_PARAGRAPH.LEFT},
        "Pelvic_Obl": {"width": Inches(3.0), "alignment": WD_ALIGN_PARAGRAPH.RIGHT},
        "HSS_Footplant": {"width": Inches(6.0), "alignment": WD_ALIGN_PARAGRAPH.CENTER},
        "Trunk_Ang_Footplant": {"width": Inches(3.0), "alignment": WD_ALIGN_PARAGRAPH.LEFT},
        "Pelvis_Ang_Footplant": {"width": Inches(3.0), "alignment": WD_ALIGN_PARAGRAPH.RIGHT},
        "Front_Leg_Brace": {"width": Inches(3.0), "alignment": WD_ALIGN_PARAGRAPH.LEFT},
        "Front_Leg_Var_Val": {"width": Inches(3.0), "alignment": WD_ALIGN_PARAGRAPH.RIGHT},
        "Lead_Leg_Midpoint": {"width": Inches(6.0), "alignment": WD_ALIGN_PARAGRAPH.CENTER},
        "Lead_Leg_GRF_z": {"width": Inches(3.0), "alignment": WD_ALIGN_PARAGRAPH.LEFT},
        "Lead_Leg_GRF_y": {"width": Inches(3.0), "alignment": WD_ALIGN_PARAGRAPH.RIGHT},
        "Horizontal_Abduction": {"width": Inches(6.0), "alignment": WD_ALIGN_PARAGRAPH.CENTER},
        "Shld_ER_Footplant": {"width": Inches(3.0), "alignment": WD_ALIGN_PARAGRAPH.LEFT},
        "Shld_ER_Max": {"width": Inches(3.0), "alignment": WD_ALIGN_PARAGRAPH.RIGHT},
        "Lateral_Trunk_Tilt": {"width": Inches(6.0), "alignment": WD_ALIGN_PARAGRAPH.CENTER},
        "Pelvis_Ang_Velo": {"width": Inches(6.0), "alignment": WD_ALIGN_PARAGRAPH.CENTER},
        "Torso_Ang_Velo": {"width": Inches(6.0), "alignment": WD_ALIGN_PARAGRAPH.CENTER},
        "Arm_Ang_Velo": {"width": Inches(6.0), "alignment": WD_ALIGN_PARAGRAPH.CENTER},
        "MPH": {"width": Inches(6.0), "alignment": WD_ALIGN_PARAGRAPH.CENTER},
    }

    # Define the list of variables to include in the report and their order
    variables_to_include = [
        "Score",
        "Linear_Pelvis_Speed",
        "Pelvic_Obl",
        "HSS_Footplant",
        "Trunk_Ang_Footplant",
        "Pelvis_Ang_Footplant",
        "Front_Leg_Brace",
        "Front_Leg_Var_Val",
        "Lead_Leg_Midpoint",
        "Lead_Leg_GRF_z",
        "Lead_Leg_GRF_y",
        "Horizontal_Abduction",
        "Shld_ER_Footplant",
        "Shld_ER_Max",
        "Lateral_Trunk_Tilt",
        "Pelvis_Ang_Velo",
        "Torso_Ang_Velo",
        "Arm_Ang_Velo",
        "MPH"
    ]
    # Connect to the SQLite reference database
    conn_reference = sqlite3.connect("grading_equation_reference_data_Pro.db")
    cursor_reference = conn_reference.cursor()

    # Iterate through variables
    for column_name in variables_to_include:
        # Add a section break
        # doc.add_section(WD_SECTION.NEW_PAGE)

        # # Add a centered title for the histogram
        # title = doc.add_paragraph()
        # title_run = title.add_run(column_name.replace('_', ' '))
        # title_run.font.size = Pt(24)
        # title_run.bold = True
        # title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add intermittent titles and descriptions
        if column_name == "Score":
            doc.add_paragraph("8ctane Score", style='Title')
            doc.add_paragraph(
                "The 8ctane score breaks down your pitch and summarizes it in a single score. This score is a product of "
                "over twenty variables that are analyzed and interpreted to determine how efficient your pitching mechanics"
                " are."
                "\n", style='Heading 2')
        elif column_name == "Linear_Pelvis_Speed":
            doc.add_paragraph("Pelvis", style='Title')
            doc.add_paragraph("Maintaining a stable and neutral pelvis throughout the pitching delivery is crucial in a "
                              "foundational sound pitch delivery."
                              "", style='Heading 2')
        elif column_name == "HSS_Footplant":
            doc.add_paragraph("Hip Shoulder Separation at FootPlant", style='Title')
            doc.add_paragraph("Hip shoulder separation is achieved by clearing the pelvis towards home while allowing the "
                              "trunk to stay as closed as possible. With pelvis angle at footplant and total hip shoulder "
                              "separation being linked to increased ball velocity, these variables are instrumental in creating "
                              "and maintaining ball velocity."
                              "", style='Heading 2')
        elif column_name == "Front_Leg_Brace":
            doc.add_paragraph("Front Leg", style='Title')
            doc.add_paragraph("Front leg brace and stability are rooted from the pelvis clearing as a cleared pelvis "
                              "restricts varus/valgus movement on the knee while enabling the front knee to extend "
                              "through ball release."
                              "", style='Heading 2')
        elif column_name == "Lead_Leg_Midpoint":
            doc.add_paragraph("Ground Reaction Forces", style='Title')
            doc.add_paragraph("Ground reaction forces are the measurement of the force the front leg puts into the ground "
                              "throughout the pitching motion. Ground reaction forces have been linked to increase ball "
                              "velocity as well as kinematic variables that are advantageous for both ball velocity and "
                              "shoulder/arm health/longevity. The ideal ground reaction force has the peaks of the ground "
                              "reaction force occurring midway between front leg footplant and ball release."
                              "", style='Heading 2')
        elif column_name == "Horizontal_Abduction":
            doc.add_paragraph("Horizontal Abduction at Footplant", style='Title')
            doc.add_paragraph("Horizontal abduction is the measurement of how far behind the body the arm is at footplant. "
                              "One of the biggest indicators of ball velocity as increased abdication allows the arm to "
                              "efficiently trail the body, allowing more force created from the front leg to transmit up "
                              "the body and act on the ball."
                              "", style='Heading 2')
        elif column_name == "Shld_ER_Footplant":
            doc.add_paragraph("Shoulder External Rotation", style='Title')
            doc.add_paragraph("Shoulder external rotation at footplant defines an early, on-time, or late arm. Late arms "
                              "are defined as below 33 degrees and an early arm is above 77 degrees at footplant. An arm "
                              "that is on-time increases the arms ability to efficiently sync up with the body while "
                              "reducing the risk of injury. "
                              "Maximum shoulder external rotation is better known as “layback”, has been correlated to "
                              "ball velocity and is an indicator of overall shoulder mobility."
                              "", style='Heading 2')
        elif column_name == "Lateral_Trunk_Tilt":
            doc.add_paragraph("Lateral Trunk Tilt at Ball Release", style='Title')
            doc.add_paragraph("Lateral trunk tilt is how tilted one is at ball release."
                              "", style='Heading 2')
        elif column_name == "Pelvis_Ang_Velo":
            doc.add_paragraph("Kinematic Sequence", style='Title')
            doc.add_paragraph("Kinematic sequencing is the order and speed at which the pelvis, torso, and arm reach "
                              "their peak rotational velocity. The order should be in sequential order of pelvis, "
                              "followed by torso, then arm. The higher the maximum rotational velocity for all segments, "
                              "the better."
                              "", style='Heading 2')
        elif column_name == "MPH":
            doc.add_paragraph("MPH", style='Title')

        # Add the histogram as an image
        filename = os.path.join(export_dir, f'{column_name}.png')
        doc.add_picture(filename, width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a new line after specific histograms
        # if column_name in ["Pelvis", "HSS_Footplant", "Trunk_Ang_Footplant", "Front_Leg_Brace", "Lead_Leg_Midpoint",
        #                    "Lead_Leg_GRF_z", "Shld_ER_Footplant", "Pelvis_Ang_Velo", "Trunk_Ang_Velo"]:
        #     doc.add_paragraph('\n')

        if column_name in ["Score", "Pelvic_Obl", "Pelvis_Ang_Footplant", "Front_Leg_Var_Val", "Lead_Leg_GRF_y",
                           "Horizontal_Abduction", "Shld_ER_Max", "Lateral_Trunk_Tilt", "Arm_Ang_Velo"]:
            doc.add_section(WD_SECTION.NEW_PAGE)

        # Set background color
        background_color = RGBColor(57, 57, 57)  # Dark gray background
        for section in doc.sections:
            section.page_color = background_color

        # font.color.theme_color =
        #
        # # Set font color
        # font_color = RGBColor(255, 255, 255)  # White font color
        # for paragraph in doc.paragraphs:
        #     for run in paragraph.runs:
        #         run.font.color.rgb = font_color

    # Close the reference database connection
    conn_reference.close()

    # Save the Word document
    doc.save(docx_filename)

except sqlite3.Error as e:
    print("Error:", e)

finally:
    # Close the connections
    conn_new.close()
    conn_reference.close()

    # Move files
    source_dir = r"C:\Users\q\PycharmProjects\Rename_Session.xml\Pitching Data\New Data"
    kinematic_data_dir = r"C:\Users\q\PycharmProjects\Rename_Session.xml\Pitching Data\Kinematic Data\Pro"
    mph_data_dir = r"C:\Users\q\PycharmProjects\Rename_Session.xml\Pitching Data\MPH Data\Pro"

    for file_name in os.listdir(source_dir):
        source_file_path = os.path.join(source_dir, file_name)
        if 'data' in file_name.lower():
            destination_file_path = os.path.join(kinematic_data_dir, file_name)
        else:
            destination_file_path = os.path.join(mph_data_dir, file_name)
        shutil.move(source_file_path, destination_file_path)
        print(f"File '{file_name}' moved to: {destination_file_path}")

    print("Files moved successfully!")

