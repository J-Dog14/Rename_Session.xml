from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import glob

def safe_add_picture(paragraph, image_path, width, height):
    """
    Attempts to add a picture to a Word document paragraph.
    If the image file does not exist, it skips adding the image.
    """
    if os.path.exists(image_path):
        paragraph.add_run().add_picture(image_path, width=width, height=height)
    else:
        print(f"Warning: Image file not found at {image_path}, skipping.")

def create_formatted_word_doc(image_paths, doc_path):
    doc = Document()

    # Footplant
    doc.add_heading('Footplant', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table = doc.add_table(rows=1, cols=2)
    cell = table.cell(0, 0)
    safe_add_picture(cell.paragraphs[0], image_paths.get('Front@FP', ''), Inches(3.21), Inches(3.21))
    cell = table.cell(0, 1)
    safe_add_picture(cell.paragraphs[0], image_paths.get('sag@FP', ''), Inches(3.21), Inches(3.21))
    doc.add_paragraph()

    # Max External Rotation and Ball Release
    heading = doc.add_heading('', level=2)
    run = heading.add_run('Max External Rotation' + ' ' * 20 + 'Ball Release')
    run.bold = True
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = doc.add_table(rows=1, cols=2)
    cell = table.cell(0, 0)
    safe_add_picture(cell.paragraphs[0], image_paths.get('sag@MaxER', ''), Inches(3.21), Inches(3.21))
    cell = table.cell(0, 1)
    safe_add_picture(cell.paragraphs[0], image_paths.get('sag@Rel', ''), Inches(3.21), Inches(3.21))
    doc.add_paragraph()

    # Kinematics Report
    doc.add_heading('Kinematics Report', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if 'KinematicReport' in image_paths:
        safe_add_picture(doc.add_paragraph(), image_paths['KinematicReport'], Inches(6), Inches(4.5))

    # Arm Path Report
    doc.add_heading('Arm Path to Release', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if 'Arm Path to Release' in image_paths:
        safe_add_picture(doc.add_paragraph(), image_paths['Arm Path to Release'], Inches(6), Inches(4.5))

    # Pronation Report
    doc.add_heading('Pronation', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if 'Pronation' in image_paths:
        safe_add_picture(doc.add_paragraph(), image_paths['Pronation'], Inches(6), Inches(4.5))
    # Save the document
    doc.save(doc_path)



# Example usage
image_paths = {
    'Front@FP': 'D:\\Youth Pronation\\Exports\\Front@FP.png',
    'sag@FP': 'D:\\Youth Pronation\\Exports\\sag@FP.png',
    'sag@MaxER': 'D:\\Youth Pronation\\Exports\\sag@MaxER.png',
    'sag@Rel': 'D:\\Youth Pronation\\Exports\\sag@Rel.png',
    'KinematicReport': 'D:\\Youth Pronation\\Exports\\KinematicReport.png',
    'Arm Path to Release': 'D:\\Youth Pronation\\Exports\\Arm Path to Release.png',
    'Pronation': 'D:\\Youth Pronation\\Exports\\Pronation.png',
}
doc_path = 'G:\\My Drive\\Action+ Reports\\New Player Data.docx'

create_formatted_word_doc(image_paths, doc_path)

def clear_images_from_folder(folder_path):
    # Define the pattern for PNG files
    pattern = os.path.join(folder_path, '*.png')

    # Use glob to find all files matching the pattern
    files = glob.glob(pattern)

    # Loop through the found files and delete them
    for file in files:
        try:
            os.remove(file)
            print(f"Deleted {file}")
        except OSError as e:
            print(f"Error: {e.strerror} - {file}")


# Example usage
source_folder = 'D:\\Youth Pronation\\Exports'
clear_images_from_folder(source_folder)
